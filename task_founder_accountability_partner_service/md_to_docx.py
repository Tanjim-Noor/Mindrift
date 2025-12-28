import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

files_map = {
    "executive_summary.md": "Executive Summary v11.docx",
    "full_report.md": "Full Report v11.docx",
    "buyer_personas_and_decision_journey.md": "Buyer Personas & Decision Journey v11.docx",
    "buyer_psychology_deep_dive.md": "Buyer Psychology Deep-Dive v11.docx",
    "pricing_psychology.md": "Pricing Psychology & Willingness-to-Pay Guidance v11.docx"
}

def add_formatted_text(paragraph, text):
    """
    Parses Markdown bold (**), italic (*), and links []() and adds them as formatted runs.
    Removes the markdown symbols.
    """
    # Pattern for bold: **text**
    # Pattern for italic: *text* (simplified, assumes no nested for this simple script)
    # Pattern for links: [text](url)
    
    # We will tokenize the string.
    # A simple way for a specific urgency is to handle Bold, then Italic, then Link.
    # But doing them sequentially on runs is distinct. 
    # Let's try a regex split approach.
    
    # Regex to capture: (**.*?**)|(\*.*?\*)|(\[.*?\]\(.*?\))
    pattern = re.compile(r'(\*\*.*?\*\*)|(\*.*?\*)|(\[.*?\]\(.*?\))')
    
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            # Extract text and url. For DOCX readable format, we'll just put Text
            # Or "Text (URL)" if needed. User wants "readable", so just Text is cleaner, maybe Text in blue.
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                link_text = m.group(1)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Skipping {md_path} (not found)")
        return

    doc = Document()
    
    # Styles config
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = None
    in_table = False

    for line in lines:
        raw_line = line.strip()
        
        # Table Handling
        if raw_line.startswith('|'):
            # Detect separator line
            if '---' in raw_line:
                continue # Skip the separator row
            
            # Parse row
            cells = [c.strip() for c in raw_line.strip('|').split('|')]
            
            if not in_table:
                # Start new table
                in_table = True
                current_table = doc.add_table(rows=0, cols=len(cells))
                current_table.style = 'Table Grid'
            
            row_cells = current_table.add_row().cells
            for i, cell_text in enumerate(cells):
                if i < len(row_cells):
                    add_formatted_text(row_cells[i].paragraphs[0], cell_text)
            continue
        else:
            in_table = False

        if not raw_line:
            continue

        # Headers
        if raw_line.startswith('# '):
            doc.add_heading(raw_line[2:], level=1)
        elif raw_line.startswith('## '):
            doc.add_heading(raw_line[3:], level=2)
        elif raw_line.startswith('### '):
            doc.add_heading(raw_line[4:], level=3)
        
        # Lists
        elif raw_line.startswith('- ') or raw_line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, raw_line[2:])
        elif re.match(r'^\d+\.\s', raw_line):
            # Ordered list
            text = re.sub(r'^\d+\.\s', '', raw_line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            
        # Blockquotes
        elif raw_line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            add_formatted_text(p, raw_line[2:])
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.italic = True
                
        # Normal Text
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, raw_line)

    doc.save(docx_path)
    print(f"Created Clean DOCX: {docx_path}")

if __name__ == "__main__":
    print("Starting smart conversion...")
    for md_file, docx_name in files_map.items():
        try:
            md_to_docx(md_file, docx_name)
        except Exception as e:
            print(f"Error converting {md_file}: {e}")
    print("Conversion complete.")
